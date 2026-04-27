"""
Simple Markdown -> DOCX converter for the lab.
- Handles headings (#..), paragraphs, bulleted lists (- ), fenced code blocks (```), and images ![alt](path).
- Images are resolved relative to the markdown file.

Usage:
    python md2docx.py lab.md lab.docx

This script requires `python-docx` (install: pip install python-docx).
"""
import sys
import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt


def add_code_block(doc, code_lines):
    p = doc.add_paragraph()
    run = p.add_run("\n".join(code_lines))
    font = run.font
    font.name = "Courier New"
    font.size = Pt(9)


def add_image(doc, img_path, max_width_in=5.0):
    try:
        doc.add_picture(str(img_path), width=Inches(max_width_in))
    except Exception as e:
        doc.add_paragraph(f"[Missing image: {img_path}]")


def convert(md_path: Path, out_path: Path):
    doc = Document()
    md_text = md_path.read_text(encoding="utf-8")
    lines = md_text.splitlines()

    in_code = False
    code_lines = []
    in_list = False

    for raw in lines:
        line = raw.rstrip('\n')

        # fenced code start/end
        if line.strip().startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                # end code block
                in_code = False
                add_code_block(doc, code_lines)
                code_lines = []
            continue

        if in_code:
            code_lines.append(line)
            continue

        # heading
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            # map # -> level 0..5 (docx heading levels start at 0 for top)
            doc.add_heading(text, level=level-1 if level-1 <= 4 else 4)
            in_list = False
            continue

        # image
        m = re.match(r"!?\[(.*?)\]\((.*?)\)", line)
        if m and line.strip().startswith("!"):
            alt = m.group(1)
            src = m.group(2)
            img_path = (md_path.parent / src).resolve()
            add_image(doc, img_path)
            in_list = False
            continue

        # bulleted list
        if line.strip().startswith("- "):
            text = line.strip()[2:].strip()
            p = doc.add_paragraph(text, style="List Bullet")
            in_list = True
            continue

        # blank line -> paragraph break
        if line.strip() == "":
            in_list = False
            continue

        # normal paragraph
        doc.add_paragraph(line)

    # if file ended while still in a code block, flush
    if in_code and code_lines:
        add_code_block(doc, code_lines)

    doc.save(str(out_path))


def main():
    if len(sys.argv) < 3:
        print("Usage: python md2docx.py input.md output.docx")
        return 2
    md = Path(sys.argv[1])
    out = Path(sys.argv[2])
    if not md.exists():
        print(f"Markdown file not found: {md}")
        return 2
    convert(md, out)
    print(f"Wrote: {out}")
    return 0


if __name__ == '__main__':
    raise SystemExit(main())
